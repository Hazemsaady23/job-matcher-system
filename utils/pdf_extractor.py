"""
PDF Extractor - Extract text from PDF files
"""

import pdfplumber
from pathlib import Path
from typing import Optional
from config.settings import MAX_PDF_PAGES

# ==========================================
# PDF EXTRACTION
# ==========================================

def extract_text_from_pdf(pdf_path: str, max_pages: int = MAX_PDF_PAGES) -> str:
    """
    Extract text from PDF file
    
    Args:
        pdf_path: Path to PDF file
        max_pages: Maximum pages to extract
    
    Returns:
        Extracted text
    """
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = min(len(pdf.pages), max_pages)
            
            for i in range(num_pages):
                page = pdf.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        return text.strip()
    
    except Exception as e:
        print(f"❌ Error extracting PDF: {e}")
        return ""

def extract_text_from_uploaded_pdf(uploaded_file) -> str:
    """
    Extract text from Streamlit uploaded PDF file
    
    Args:
        uploaded_file: Streamlit UploadedFile object
    
    Returns:
        Extracted text
    """
    try:
        text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages[:MAX_PDF_PAGES]:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        return text.strip()
    
    except Exception as e:
        print(f"❌ Error extracting uploaded PDF: {e}")
        return ""

# ==========================================
# DOCX EXTRACTION
# ==========================================

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from DOCX file
    
    Args:
        docx_path: Path to DOCX file
    
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        
        doc = Document(docx_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        return "\n".join(text)
    
    except Exception as e:
        print(f"❌ Error extracting DOCX: {e}")
        return ""

def extract_text_from_uploaded_docx(uploaded_file) -> str:
    """
    Extract text from Streamlit uploaded DOCX file
    
    Args:
        uploaded_file: Streamlit UploadedFile object
    
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        
        doc = Document(uploaded_file)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        return "\n".join(text)
    
    except Exception as e:
        print(f"❌ Error extracting uploaded DOCX: {e}")
        return ""

# ==========================================
# GENERIC FILE HANDLER
# ==========================================

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from file (auto-detect type)
    
    Args:
        file_path: Path to file
    
    Returns:
        Extracted text
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(str(file_path))
    elif extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {extension}")

def extract_text_from_uploaded_file(uploaded_file):
    """
    Extract text from Streamlit uploaded file (auto-detect type)
    
    Args:
        uploaded_file: Streamlit UploadedFile object
    
    Returns:
        Extracted text
    """
    file_type = uploaded_file.type
    
    if file_type == 'application/pdf':
        return extract_text_from_uploaded_pdf(uploaded_file)
    elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        return extract_text_from_uploaded_docx(uploaded_file)
    elif file_type == 'text/plain':
        return uploaded_file.read().decode('utf-8')
    else:
        # Try by filename extension
        if uploaded_file.name.endswith('.pdf'):
            return extract_text_from_uploaded_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            return extract_text_from_uploaded_docx(uploaded_file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

# ==========================================
# PDF INFO
# ==========================================

def get_pdf_info(pdf_path: str) -> dict:
    """Get PDF metadata"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            return {
                'num_pages': len(pdf.pages),
                'metadata': pdf.metadata
            }
    except Exception as e:
        return {'error': str(e)}

# ==========================================
# TESTING
# ==========================================

if __name__ == "__main__":
    print("✅ PDF Extractor ready!")
    print("\nFunctions available:")
    print("  - extract_text_from_pdf()")
    print("  - extract_text_from_docx()")
    print("  - extract_text_from_file()")
    print("  - extract_text_from_uploaded_file()")